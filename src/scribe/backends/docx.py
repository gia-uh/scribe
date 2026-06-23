from __future__ import annotations

import io

import docx

from ..result import ExtractResult


def _table_md(table) -> str:
    rows = [[c.text.strip().replace("\n", " ") for c in r.cells] for r in table.rows]
    if not rows:
        return ""
    head, *body = rows
    out = ["| " + " | ".join(head) + " |", "| " + " | ".join(["---"] * len(head)) + " |"]
    for r in body:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def convert(data: bytes) -> ExtractResult:
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"could not open DOCX: {exc}") from exc

    parts: list[str] = []
    title: str | None = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            level = min(max(level, 1), 6)
            if title is None and level == 1:
                title = text
            parts.append("#" * level + " " + text)
        elif "list bullet" in style:
            parts.append("- " + text)
        elif "list number" in style:
            parts.append("1. " + text)
        else:
            parts.append(text)

    for table in doc.tables:
        md = _table_md(table)
        if md:
            parts.append(md)

    return ExtractResult(markdown="\n\n".join(parts), title=title, meta={"backend": "docx"})
