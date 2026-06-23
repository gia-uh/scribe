import io

import docx

from scribe.backends.docx import convert


def _make():
    d = docx.Document()
    d.add_heading("Title One", level=1)
    d.add_heading("Sub", level=2)
    d.add_paragraph("Body paragraph.")
    d.add_paragraph("first item", style="List Bullet")
    d.add_paragraph("second item", style="List Bullet")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "A"
    t.rows[0].cells[1].text = "B"
    t.rows[1].cells[0].text = "1"
    t.rows[1].cells[1].text = "2"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_structure():
    md = convert(_make()).markdown
    assert "# Title One" in md
    assert "## Sub" in md
    assert "- first item" in md and "- second item" in md
    assert "| A | B |" in md


def test_docx_title_captured():
    assert convert(_make()).title == "Title One"


def test_docx_corrupt_raises():
    import pytest

    with pytest.raises(ValueError):
        convert(b"not a docx")
